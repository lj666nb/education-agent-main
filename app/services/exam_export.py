"""试卷导出：PDF (fpdf2) + Word (python-docx)"""
import re
from io import BytesIO
from typing import List, Dict, Any, Optional
from uuid import UUID
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session
from app.models.question_bank import Question

QTYPE_LABELS = {
    "single_choice": "单选题", "multiple_choice": "多选题",
    "fill_blank": "填空题", "true_false": "判断题",
    "short_answer": "简答题", "programming": "编程题", "essay": "论述题",
}


# ── LaTeX math to Unicode conversion ──

_LATEX_UNICODE = {
    # Relations
    r"\\subseteq": "⊆", r"\\supseteq": "⊇",
    r"\\subset": "⊂", r"\\supset": "⊃",
    r"\\in": "∈", r"\\notin": "∉",
    r"\\ni": "∋", r"\\not\\ni": "∌",
    r"\\equiv": "≡", r"\\cong": "≅",
    r"\\approx": "≈", r"\\sim": "∼",
    r"\\neq": "≠", r"\\leq": "≤", r"\\geq": "≥",
    r"\\ll": "≪", r"\\gg": "≫",
    r"\\perp": "⊥", r"\\parallel": "∥",
    r"\\mid": "∣", r"\\propto": "∝",
    r"\\models": "⊨",
    r"\\vdash": "⊢", r"\\dashv": "⊣",
    # Arrows
    r"\\rightarrow": "→", r"\\leftarrow": "←",
    r"\\leftrightarrow": "↔",
    r"\\Rightarrow": "⇒", r"\\Leftarrow": "⇐",
    r"\\Leftrightarrow": "⇔",
    r"\\longrightarrow": "⟶", r"\\longleftarrow": "⟵",
    r"\\Longrightarrow": "⟹", r"\\Longleftarrow": "⟸",
    r"\\Longleftrightarrow": "⟺",
    r"\\mapsto": "↦", r"\\longmapsto": "⟼",
    r"\\to": "→", r"\\gets": "←",
    r"\\uparrow": "↑", r"\\downarrow": "↓",
    r"\\Uparrow": "⇑", r"\\Downarrow": "⇓",
    # Operators
    r"\\cup": "∪", r"\\cap": "∩",
    r"\\wedge": "∧", r"\\vee": "∨",
    r"\\oplus": "⊕", r"\\otimes": "⊗",
    r"\\odot": "⊙", r"\\circ": "∘",
    r"\\bullet": "∙", r"\\cdot": "⋅",
    r"\\times": "×", r"\\div": "÷",
    r"\\pm": "±", r"\\mp": "∓",
    r"\\ast": "∗", r"\\star": "⋆",
    r"\\lnot": "¬", r"\\neg": "¬",
    r"\\land": "∧", r"\\lor": "∨",
    r"\\oplus": "⊕", r"\\ominus": "⊖",
    r"\\sqcup": "⊔", r"\\sqcap": "⊓",
    r"\\nabla": "∇", r"\\partial": "∂",
    r"\\prime": "′", r"\\emptyset": "∅",
    r"\\varnothing": "∅",
    # Quantifiers
    r"\\forall": "∀", r"\\exists": "∃",
    r"\\nexists": "∄",
    # Sets
    r"\\mathbb\{N\}": "ℕ", r"\\mathbb\{Z\}": "ℤ",
    r"\\mathbb\{Q\}": "ℚ", r"\\mathbb\{R\}": "ℝ",
    r"\\mathbb\{C\}": "ℂ",
    # Greek lowercase
    r"\\alpha": "α", r"\\beta": "β",
    r"\\gamma": "γ", r"\\delta": "δ",
    r"\\epsilon": "ε", r"\\varepsilon": "ε",
    r"\\zeta": "ζ", r"\\eta": "η",
    r"\\theta": "θ", r"\\vartheta": "ϑ",
    r"\\iota": "ι", r"\\kappa": "κ",
    r"\\lambda": "λ", r"\\mu": "μ",
    r"\\nu": "ν", r"\\xi": "ξ",
    r"\\omicron": "ο", r"\\pi": "π",
    r"\\varpi": "ϖ", r"\\rho": "ρ",
    r"\\varrho": "ϱ", r"\\sigma": "σ",
    r"\\varsigma": "ς", r"\\tau": "τ",
    r"\\upsilon": "υ", r"\\phi": "φ",
    r"\\varphi": "φ", r"\\chi": "χ",
    r"\\psi": "ψ", r"\\omega": "ω",
    # Greek uppercase
    r"\\Gamma": "Γ", r"\\Delta": "Δ",
    r"\\Theta": "Θ", r"\\Lambda": "Λ",
    r"\\Xi": "Ξ", r"\\Pi": "Π",
    r"\\Sigma": "Σ", r"\\Phi": "Φ",
    r"\\Psi": "Ψ", r"\\Omega": "Ω",
    # Miscellaneous
    r"\\infty": "∞",
    r"\\angle": "∠", r"\\measuredangle": "∡",
    r"\\triangle": "△", r"\\Box": "□",
    r"\\diamond": "◇", r"\\star": "⋆",
    r"\\ell": "ℓ", r"\\hbar": "ℏ",
    r"\\hslash": "ℏ", r"\\Im": "ℑ",
    r"\\Re": "ℜ", r"\\wp": "℘",
    r"\\aleph": "ℵ", r"\\beth": "ℶ",
    r"\\dag": "†", r"\\ddag": "‡",
    r"\\S": "§", r"\\P": "¶",
    r"\\dots": "…", r"\\cdots": "⋯",
    r"\\vdots": "⋮", r"\\ddots": "⋱",
    # Function names (keep as text)
    r"\\sin": "sin", r"\\cos": "cos", r"\\tan": "tan",
    r"\\log": "log", r"\\ln": "ln", r"\\lg": "lg",
    r"\\max": "max", r"\\min": "min",
    r"\\lim": "lim", r"\\sup": "sup", r"\\inf": "inf",
    r"\\det": "det", r"\\deg": "deg",
    r"\\mod": "mod", r"\\bmod": "mod", r"\\pmod": "(mod",
}


def _latex_to_unicode(text: str) -> str:
    """Convert LaTeX math markup to Unicode symbols."""
    if not text:
        return text

    # Remove inline math delimiters
    text = re.sub(r"\$\$(.*?)\$\$", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\$(.*?)\$", r"\1", text)

    # Replace known LaTeX commands (longer patterns first to avoid partial matches)
    for cmd, unichr in sorted(_LATEX_UNICODE.items(), key=lambda x: -len(x[0])):
        text = text.replace(cmd, unichr)

    # Handle superscripts: ^{...} and ^char
    def _superscript(m):
        content = m.group(1) or m.group(2)
        return "".join(_SUP_MAP.get(c, c) for c in content)
    text = re.sub(r"\^\{([^}]+)\}", _superscript, text)
    text = re.sub(r"\^([a-zA-Z0-9])", lambda m: _SUP_MAP.get(m.group(1), m.group(1)), text)

    # Handle subscripts: _{...} and _char
    def _subscript(m):
        content = m.group(1) or m.group(2)
        return "".join(_SUB_MAP.get(c, c) for c in content)
    text = re.sub(r"\_\{([^}]+)\}", _subscript, text)
    text = re.sub(r"_([a-zA-Z0-9])", lambda m: _SUB_MAP.get(m.group(1), m.group(1)), text)

    # Clean up extra spaces
    text = re.sub(r"\s+", " ", text).strip()
    return text


_SUP_MAP = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³",
    "4": "⁴", "5": "⁵", "6": "⁶", "7": "⁷",
    "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼",
    "(": "⁽", ")": "⁾",
    "n": "ⁿ",
    "*": "⋆",
}
_SUB_MAP = {
    "0": "₀", "1": "₁", "2": "₂", "3": "₃",
    "4": "₄", "5": "₅", "6": "₆", "7": "₇",
    "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌",
    "(": "₍", ")": "₎",
}


def _convert_question_text(content: Any) -> Any:
    """Recursively convert LaTeX in question content/answer fields."""
    if isinstance(content, str):
        return _latex_to_unicode(content)
    if isinstance(content, dict):
        return {k: _convert_question_text(v) for k, v in content.items()}
    if isinstance(content, list):
        return [_convert_question_text(v) for v in content]
    return content


# ── PDF export ──

class ExamPDF(FPDF):
    def __init__(self, font_name="DejaVu"):
        super().__init__()
        self._exam_font_name = font_name

    def header(self):
        if self.page_no() > 1:
            self.set_font(self._exam_font_name, "", 8)
            self.cell(0, 8, "试卷", align="C", new_x="LMARGIN", new_y="NEXT")

    def footer(self):
        self.set_y(-15)
        self.set_font(self._exam_font_name, "", 8)
        self.cell(0, 10, f"第 {self.page_no()} 页", align="C")


def export_pdf(
    title: str,
    sections: List[Dict[str, Any]],
    questions_by_id: Dict[str, Question],
    total_score: int,
    time_limit_minutes: Optional[int] = None,
) -> BytesIO:
    """生成 PDF 试卷"""

    import os
    import atexit

    # ── 字体选择策略 ──
    # 优先使用 TrueType 格式的中文字体（CIDFontType2），pymupdf 渲染兼容性更好。
    # 从 TTC 中提取 TTF 字体文件以绕过 fpdf2 的 CFF CID 字体 (CIDFontType0C) 渲染问题。
    _TTF_CACHE: Dict[str, str] = {}

    def _extract_ttf_from_ttc(ttc_path: str, font_index: int, cache_key: str) -> Optional[str]:
        """从 TTC 字体集合中提取单个 TTF 字体文件并缓存"""
        if cache_key in _TTF_CACHE and os.path.exists(_TTF_CACHE[cache_key]):
            return _TTF_CACHE[cache_key]
        try:
            from fontTools.ttLib import TTCollection
            ttc = TTCollection(ttc_path)
            if font_index < len(ttc):
                ttf_path = f"/tmp/{cache_key}.ttf"
                ttc[font_index].save(ttf_path)
                _TTF_CACHE[cache_key] = ttf_path
                atexit.register(lambda p=ttf_path: os.path.exists(p) and os.remove(p))
                return ttf_path
        except Exception:
            pass
        return None

    # 字体候选列表：(label, regular_path, bold_path, font_name)
    _FONT_CANDIDATES = []

    # 1) WenQuanYi Micro Hei（TrueType，优先，兼容性最佳）
    _WQY_TTC = "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"
    if os.path.exists(_WQY_TTC):
        wqy_regular = _extract_ttf_from_ttc(_WQY_TTC, 0, "wqy_microhei_regular")
        wqy_bold = wqy_regular  # WQY Micro Hei 只有 Regular 和 Mono，无独立 Bold
        if wqy_regular:
            _FONT_CANDIDATES.append(("WQY", wqy_regular, wqy_bold, "WenQuanYiMicroHei"))

    # 2) Noto Sans CJK SC（CFF CID，回退方案；需要提取 TTF 避免 CIDFontType0C）
    _NOTO_TTC = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
    _NOTO_BOLD_TTC = "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
    if os.path.exists(_NOTO_TTC):
        noto_regular = _extract_ttf_from_ttc(_NOTO_TTC, 2, "notocjk_regular_sc")  # SC = Simplified Chinese
        noto_bold = None
        if os.path.exists(_NOTO_BOLD_TTC):
            noto_bold = _extract_ttf_from_ttc(_NOTO_BOLD_TTC, 2, "notocjk_bold_sc")
        if noto_regular:
            _FONT_CANDIDATES.append(("NotoSC", noto_regular, noto_bold or noto_regular, "NotoSansCJKsc"))

    # 3) DejaVu Sans（仅支持拉丁字符，最后回退）
    _DEJAVU_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    _DEJAVU_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    if os.path.exists(_DEJAVU_REGULAR):
        bold = _DEJAVU_BOLD if os.path.exists(_DEJAVU_BOLD) else _DEJAVU_REGULAR
        _FONT_CANDIDATES.append(("DejaVu", _DEJAVU_REGULAR, bold, "DejaVu"))

    if not _FONT_CANDIDATES:
        raise RuntimeError("未找到字体文件，请安装 fonts-wqy-microhei 或 fonts-noto-cjk 或 fonts-dejavu-core")

    label, regular_font, bold_font, font_name = _FONT_CANDIDATES[0]

    pdf = ExamPDF(font_name=font_name)
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.add_font(font_name, "", regular_font)
    pdf.add_font(font_name, "B", bold_font)

    pdf.add_page()
    pdf.set_font(font_name, "B", 18)
    pdf.cell(0, 15, title, align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font(font_name, size=11)
    info = [f"总分：{total_score} 分", f"总题数：{sum(s.get('count', 0) for s in sections)} 题"]
    if time_limit_minutes:
        info.append(f"时间：{time_limit_minutes} 分钟")
    pdf.cell(0, 10, "  |  ".join(info), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)

    q_num = 1
    answer_key = []
    for section in sections:
        pdf.set_font(font_name, "B", 13)
        sec_info = f"{section.get('name', '')}（共{section.get('count', 0)}题，每题{section.get('score_per_question', 0)}分）"
        pdf.cell(0, 12, sec_info, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        for qid in section.get("question_ids", []):
            q = questions_by_id.get(qid)
            if not q:
                continue
            stem = _convert_question_text(q.content.get("stem", ""))
            pdf.set_font(font_name, size=11)
            pdf.multi_cell(0, 7, f"{q_num}. {stem}", new_x="LMARGIN", new_y="NEXT")

            for opt in q.content.get("options", []):
                pdf.set_font(font_name, size=10)
                opt_text = _convert_question_text(opt['text'])
                pdf.cell(0, 7, f"    {opt['key']}. {opt_text}", new_x="LMARGIN", new_y="NEXT")

            # 填空/简答留空
            if q.type in ("fill_blank", "short_answer", "essay"):
                pdf.ln(10)

            pdf.ln(3)
            answer_key.append({
                "num": q_num,
                "answer": q.answer.get("correct_answer", []),
                "explanation": q.answer.get("explanation", ""),
            })
            q_num += 1

    # 答案页
    pdf.add_page()
    pdf.set_font(font_name, "B", 16)
    pdf.cell(0, 15, "参考答案", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    for item in answer_key:
        ans = ", ".join(item["answer"]) if item["answer"] else "略"
        ans = _convert_question_text(ans)
        pdf.set_font(font_name, size=11)
        pdf.cell(0, 7, f"{item['num']}. {ans}", new_x="LMARGIN", new_y="NEXT")
        if item.get("explanation"):
            pdf.set_font(font_name, size=9)
            expl = _convert_question_text(item["explanation"])
            pdf.multi_cell(0, 5, f"    解析：{expl}")

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


# ── Word export ──

def export_word(
    title: str,
    sections: List[Dict[str, Any]],
    questions_by_id: Dict[str, Question],
    total_score: int,
    time_limit_minutes: Optional[int] = None,
) -> BytesIO:
    """生成 Word 试卷"""
    doc = Document()

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    # 信息行
    info = [f"总分：{total_score} 分"]
    if time_limit_minutes:
        info.append(f"时间：{time_limit_minutes} 分钟")
    info.append(f"总题数：{sum(s.get('count', 0) for s in sections)} 题")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("  |  ".join(info))
    run.font.size = Pt(11)
    doc.add_paragraph()

    q_num = 1
    answer_items = []
    for section in sections:
        p = doc.add_paragraph()
        run = p.add_run(
            f"{section.get('name', '')}（共{section.get('count', 0)}题，每题{section.get('score_per_question', 0)}分）")
        run.bold = True
        run.font.size = Pt(13)

        for qid in section.get("question_ids", []):
            q = questions_by_id.get(qid)
            if not q:
                continue
            stem = _convert_question_text(q.content.get("stem", ""))
            doc.add_paragraph(f"{q_num}. {stem}")
            for opt in q.content.get("options", []):
                opt_text = _convert_question_text(opt['text'])
                doc.add_paragraph(f"    {opt['key']}. {opt_text}")
            if q.type in ("fill_blank", "short_answer", "essay"):
                doc.add_paragraph()
            doc.add_paragraph()
            answer_items.append({
                "num": q_num,
                "answer": q.answer.get("correct_answer", []),
            })
            q_num += 1

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("参考答案")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()
    for item in answer_items:
        ans = ", ".join(item["answer"]) if item["answer"] else "略"
        doc.add_paragraph(f"{item['num']}. {ans}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
